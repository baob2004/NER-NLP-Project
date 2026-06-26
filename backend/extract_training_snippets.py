from __future__ import annotations

import argparse
import re
import unicodedata
from pathlib import Path
from typing import List, Optional, Tuple

import fitz
from docx import Document
from docx.shared import Pt


DEFAULT_INPUT_DIR = Path("../data/raw_judgments")
DEFAULT_OUTPUT_DIR = Path("../data/snippets_raw")


def _nfc(text: str) -> str:
    text = unicodedata.normalize("NFC", text)
    for bad, good in [
        ("\u00a0", " "),
        ("\u200b", ""),
        ("\u200c", ""),
        ("\u200d", ""),
        ("\ufeff", ""),
        ("\ufffe", " "),
        ("\u2013", "-"),
        ("\u2014", "-"),
        ("\u201c", '"'),
        ("\u201d", '"'),
        ("\u2018", "'"),
        ("\u2019", "'"),
    ]:
        text = text.replace(bad, good)

    fixes = {
        "bịcáo": "bị cáo",
        "Tuyênbố": "Tuyên bố",
        "Xửphạt": "Xử phạt",
        "chấtma": "chất ma",
        "matúy": "ma túy",
        "ma tuý": "ma túy",
        "tàngtrữ": "tàng trữ",
        "tráiphép": "trái phép",
        "Bộluật": "Bộ luật",
        "Hìnhsự": "Hình sự",
    }
    for a, b in fixes.items():
        text = text.replace(a, b)

    text = re.sub(r"[ \t]+", " ", text)
    return text


def _strip_diacritics(text: str) -> str:
    text = unicodedata.normalize("NFC", text)
    text = text.replace("Đ", "D").replace("đ", "d")
    nfd = unicodedata.normalize("NFD", text)
    return "".join(c for c in nfd if unicodedata.category(c) != "Mn").upper()


def clean_raw(text: str) -> str:
    text = _nfc(text)
    lines = []

    for line in text.splitlines():
        s = line.strip()

        if re.fullmatch(r"\d{1,3}", s):
            continue

        if re.fullmatch(r"[-–—_\s]{3,}", s):
            continue

        lines.append(line.rstrip())

    text = "\n".join(lines)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_pdf(pdf_path: Path) -> str:
    doc = fitz.open(str(pdf_path))
    pages = [page.get_text("text", sort=True) for page in doc]
    return clean_raw("\n".join(pages))


_ND_HEADS = [
    r"NHAN\s+DINH\s+CUA\s+TOA\s+AN",
    r"NHAN\s+DINH\s+CUA\s+HOI\s+DONG\s+XET\s+XU",
    r"HOI\s+DONG\s+XET\s+XU\s+NHAN\s+DINH",
    r"TOA\s+AN\s+NHAN\s+DINH",
    r"XET\s+THAY\s*:",
]

_CONTENT_HEADS = [
    r"NOI\s+DUNG\s+VU\s+AN\s*:?",
]

_QD_HEADS = [
    r"(?m)^\s*QUYET\s+DINH\s*:?\s*$",
    r"(?m)^\s*QUYET\s+DINH\s*:",
]

_QD_STOP = [
    r"(?<![,\w\d])\s*2\.\s+(VE\s+)?VIEC\s+XU\s+LY\s+VAT\s+CHUNG",
    r"(?<![,\w\d])\s*2\.\s+XU\s+LY\s+VAT\s+CHUNG",
    r"(?<![,\w\d])\s*2\.\s+BIEN\s+PHAP\s+TU\s+PHAP",
    r"(?<![,\w\d])\s*[23]\.\s+(VE\s+)?AN\s+PHI",
    r"(?<![,\w\d])\s*[34]\.\s+(VE\s+)?QUYEN\s+KHANG\s+CAO",
    r"\bVE\s+VAT\s+CHUNG\s+VU\s+AN\b",
    r"\bTICH\s+THU\s+TIEU\s+HUY\b",
    r"\bTICH\s+THU\s+SUNG\s+(VAO\s+)?NGAN\s+SACH\b",
    r"TM\.\s*HOI\s+DONG\s+XET\s+XU",
    r"NOI\s+NHAN\s*:",
    r"THAM\s+PHAN.{0,30}CHU\s+TOA",
    r"\(DA\s+KY\)",
    r"TRUONG\s+HOP\s+BAN\s+AN\s+DUOC\s+THI\s+HANH",
    r"LUAT\s+THI\s+HANH\s+AN\s+DAN\s+SU",
]


def _first_pos(norm_text: str, patterns: list[str], start: int = 0) -> Optional[int]:
    best = None
    for pat in patterns:
        m = re.search(pat, norm_text[start:], re.IGNORECASE | re.MULTILINE)
        if m:
            pos = start + m.start()
            if best is None or pos < best:
                best = pos
    return best


def _locate_sections(text: str) -> Tuple[Optional[int], Optional[int], Optional[int]]:
    norm = _strip_diacritics(text)
    content = _first_pos(norm, _CONTENT_HEADS)
    nd = _first_pos(norm, _ND_HEADS)

    qd_start_search = (nd + 200) if nd is not None else 0
    qd = _first_pos(norm, _QD_HEADS, start=qd_start_search)

    return content, nd, qd


_DRUG_NAME_PAT = re.compile(r"""
    \b(?:
      [Mm]ethamphetamine | [Mm]etamphetamine | [Mm]etamphetamin | [Mm]ethamphetamin | [Mm]ethmphetamine
    | [Aa]mphetamine     | [Aa]mphetamin
    | [Hh]eroin[e]?
    | [Kk]etamine        | [Kk]etamin
    | MDMA | MDA | MDEA
    | [Mm]orphine        | [Mm]orphin
    | [Cc]ocaine         | [Cc]ocain
    | [Cc]odeine         | [Cc]odein
    | [Mm]ethadone
    | [Ff]entanyl
    | LSD | PCP | GHB
    | [Pp]silocybin      | [Pp]silocin
    | [Oo]pium
    | [Cc]annabis        | [Mm]arijuana

    # Tiếng Việt có dấu
    | [Tt]huốc\s+phiện   | [Aa]\s+phiện
    | [Cc]ần\s+sa
    | [Cc]ỏ\s+[Mm]ỹ
    | [Tt]huốc\s+lắc     | [Kk]ẹo\s+lắc | [Mm]a\s+túy\s+kẹo
    | [Mm]a\s*túy\s+đá   | [Mm]a\s*tuý\s+đá | [Hh]àng\s+đá
    | [Hh]ồng\s+phiến
    | [Nn]ước\s+vui      | [Hh]appy\s+[Ww]ater
    | [Nn]ước\s+xoài     | [Nn]ước\s+nho
    | [Tt]em\s+giấy      | [Bb]ùa\s+lưỡi
    | [Mm]a\s*túy\s+khay | [Hh]àng\s+khay
    | [Nn]ấm\s+thức\s+thần
    | [Tt]hảo\s+mộc\s+khô

    # Tiếng Việt không dấu/font lỗi
    | ma\s*tuy\s+da      | hang\s+da
    | thuoc\s+phien      | a\s+phien
    | can\s+sa           | co\s+my
    | thuoc\s+lac        | keo\s+lac
    | hong\s+phien
    | nuoc\s+vui         | nuoc\s+xoai
    | tem\s+giay         | bua\s+luoi
    | thao\s+moc\s+kho
    )\b
""", re.VERBOSE | re.IGNORECASE)

_WEIGHT_PAT = re.compile(
    r"\b(\d+(?:[.,]\d+)?)\s*(gam|gram|g|kg|kilogam|miligam|milligam|mg)\b",
    re.IGNORECASE,
)

_RANGE_BEFORE = re.compile(r"\b(Từ|Dưới|trên)\s*$", re.IGNORECASE)
_RANGE_AFTER = re.compile(r"^\s*(?:đến|đến dưới|trở lên|trở xuống)\b", re.IGNORECASE)

_CRIME_ACTION_PAT = re.compile(r"""
    \b(?:
      phạm\s+tội
    | tàng\s+trữ
    | cất\s+giấu
    | cất\s+giữ
    | cất\s+trên\s+người
    | mua\s+bán
    | bán\s+trái\s+phép
    | mua\s+.*ma\s+túy
    | mua\s+.*ma\s+tuý
    | bán\s+.*ma\s+túy
    | bán\s+.*ma\s+tuý
    | vận\s+chuyển
    | sử\s+dụng
    | tổ\s+chức\s+sử\s+dụng
    | hùn\s+tiền
    | đặt\s+mua
    | giao\s+ma\s+túy
    | giao\s+ma\s+tuý
    | nhận\s+ma\s+túy
    | nhận\s+ma\s+tuý
    | giao\s+nhận
    | mang\s+theo
    | mang\s+về
    | bắt\s+quả\s+tang
    | bị\s+bắt\s+quả\s+tang
    | kiểm\s+tra
    | phát\s+hiện
    | thu\s+giữ
    | khám\s+xét
    )\b
""", re.VERBOSE | re.IGNORECASE)

_TIME_SIGNAL_PAT = re.compile(r"""
    \b(?:
      Khoảng
    | Vào\s+khoảng
    | Vào\s+lúc
    | Hồi
    | Đến\s+khoảng
    | Vào\s+hồi
    | Ngày
    | Chiều\s+ngày
    | Tối\s+ngày
    | Sáng\s+ngày
    | Trưa\s+ngày
    | Rạng\s+sáng
    | Từ\s+ngày
    )\b
    |
    \b\d{1,2}\s*(?:giờ|h)\s*(?:\d{1,2}\s*phút)?\b
    |
    \b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b
""", re.VERBOSE | re.IGNORECASE)

_LOC_SIGNAL_PAT = re.compile(r"""
    \b(?:
      tại
    | ở
    | khu\s+vực
    | trước\s+số\s+nhà
    | phòng\s+trọ
    | nhà\s+trọ
    | khách\s+sạn
    | nhà\s+nghỉ
    | quán
    | đường
    | phố
    | thôn
    | xóm
    | ấp
    | tổ
    | khu\s+phố
    | xã
    | phường
    | thị\s+trấn
    | huyện
    | quận
    | thành\s+phố
    | tỉnh
    )\b
""", re.VERBOSE | re.IGNORECASE)


def _has_drug_name(text: str) -> bool:
    return bool(_DRUG_NAME_PAT.search(text))


def _has_real_weight(text: str) -> bool:
    for m in _WEIGHT_PAT.finditer(text):
        before = text[max(0, m.start() - 25):m.start()]
        after = text[m.end():m.end() + 40]
        is_range = (
            bool(_RANGE_BEFORE.search(before)) and bool(_RANGE_AFTER.search(after))
        ) or bool(_RANGE_AFTER.search(after))

        if is_range and not _has_drug_name(text):
            continue

        return True
    return False


def _has_crime_action(text: str) -> bool:
    return bool(_CRIME_ACTION_PAT.search(text))


def _has_time_signal(text: str) -> bool:
    return bool(_TIME_SIGNAL_PAT.search(text))


def _has_loc_signal(text: str) -> bool:
    return bool(_LOC_SIGNAL_PAT.search(text))


def _has_drug_or_weight(text: str) -> bool:
    return _has_drug_name(text) or _has_real_weight(text)


_NOISE_HEAD_PAT = re.compile(r"""
    ^\s*\[\d+\]\s*
    (?:
      Ve\s+hanh\s+vi.*quyet\s+dinh\s+to\s+tung
    | Ve\s+to\s+tung
    | Ve\s+thu\s+tuc
    | Cac?\s+quyet\s+dinh.*co\s+quan\s+tien\s+hanh
    | Hanh\s+vi.*quyet\s+dinh.*to\s+tung
    | Nguoi\s+lam\s+chung
    | Ve\s+nguoi\s+lam\s+chung
    | Viec\s+vang\s+mat
    | Ve\s+loi
    | Loi\s+cua\s+bi\s+cao
    | Ve\s+tinh\s+tiet\s+tang\s+nang
    | Ve\s+tinh\s+tiet.*giam\s+nhe
    | Tinh\s+tiet\s+tang\s+nang
    | Tinh\s+tiet.*giam\s+nhe
    | Ve\s+nhan\s+than
    | Nhan\s+than\s+bi\s+cao
    | Ve\s+hinh\s+phat\s+bo\s+sung
    | Hinh\s+phat\s+bo\s+sung
    | Ve\s+muc\s+hinh\s+phat
    | Ve\s+hinh\s+phat(?!\s+chinh)
    | Muc\s+hinh\s+phat
    | (?:Ve\s+)?(?:xu\s+ly\s+)?vat\s+chung
    | (?:Ve\s+)?an\s+phi
    )
""", re.VERBOSE | re.IGNORECASE)

_NOISE_INLINE_PAT = re.compile(r"""
    NGAY\s+XET\s+XU
  | XET\s+XU\s+SO\s+THAM
  | THU\s+LY
  | CAO\s+TRANG
  | QUYET\s+DINH\s+DUA\s+VU\s+AN
  | TRU\s+SO\s+TOA\s+AN
  | HOI\s+TRUONG\s+XET\s+XU
  | KIEM\s+SAT\s+VIEN
  | LUAT\s+SU
  | NGUOI\s+BAO\s+CHUA
  | THAM\s+PHAN
  | HOI\s+THAM
  | THU\s+KY
  | HKTT
  | NOI\s+CU\s+TRU
  | DIA\s+CHI\s+CU\s+TRU
  | CON\s+ONG
  | CON\s+BA
  | TIEN\s+AN
  | TIEN\s+SU
  | TAM\s+GIU
  | TAM\s+GIAM
  | THOI\s+HAN\s+TU\s+TINH\s+TU
  | MAU\s+NUOC\s+TIEU
  | NUOC\s+TIEU
  | TEST\s+NHANH
  | XET\s+NGHIEM\s+DUONG\s+TINH
  | TINH\s+TIET\s+TANG\s+NANG
  | TINH\s+TIET\s+GIAM\s+NHE
  | HINH\s+PHAT\s+BO\s+SUNG
  | HANH\s+VI\s+NGUY\s+HIEM\s+CHO\s+XA\s+HOI
  | CHINH\s+SACH\s+DOC\s+QUYEN
  | MUC\s+DO\s+PHAM\s+TOI
  | AN\s+PHI
  | QUYEN\s+KHANG\s+CAO
  | NOI\s+NHAN
  | TICH\s+THU\s+TIEU\s+HUY
  | TICH\s+THU\s+SUNG\s+NGAN\s+SACH
  | \bDIEU\s+106\b
  | NGHI\s+QUYET\s+326
  | HOAN\s+LAI\s+SAU\s+GIAM\s+DINH
  | CON\s+LAI\s+SAU\s+GIAM\s+DINH
  | MAU\s+VAT\s+CON\s+LAI
  | MAU\s+HOAN\s+LAI
""", re.VERBOSE | re.IGNORECASE)


def _is_noise(text: str) -> bool:
    n = _strip_diacritics(text)
    head_120 = _strip_diacritics(text[:120])
    return bool(_NOISE_HEAD_PAT.search(head_120)) or bool(_NOISE_INLINE_PAT.search(n))


_SECTION_TITLE_RE = re.compile(
    r"^\s*(?:NỘI\s+DUNG\s+VỤ\s+ÁN|NHẬN\s+ĐỊNH\s+CỦA\s+TÒA\s+ÁN"
    r"|NHẬN\s+ĐỊNH\s+CỦA\s+HỘI\s+ĐỒNG\s+XÉT\s+XỬ"
    r"|XÉT\s+THẤY|QUYẾT\s+ĐỊNH)\s*:?\s*$",
    re.IGNORECASE,
)

_SENT_SPLIT = re.compile(
    r"(?<=[.!?;:])\s+(?=[A-ZĐÁÀẢÃẠĂẮẰẶÂẤẦẬÉÈỆÊẾỀÍÌỊÓÒỌÔỐỒỘƠỚỜỢÚÙỤƯỨỪỰÝỲỴ\[])"
)


def _strip_tail_noise(unit: str) -> str:
    patterns = [
        r"\s*\*\s*Vật\s+chứng\s+thu\s+giữ\s*:.*$",
        r"\s*Vật\s+chứng\s+thu\s+giữ\s*:.*$",
        r"\s*Biện\s+pháp\s+tư\s+pháp\s*:.*$",
        r"\s*Về\s+vật\s+chứng\s*:.*$",
    ]
    out = unit
    for pat in patterns:
        out = re.sub(pat, "", out, flags=re.IGNORECASE | re.DOTALL).strip()
    return out


def _to_units(text: str) -> List[str]:
    text = re.sub(r"(?<!\A)\s*(\[\d+\]\s*)", r"\n\n\1", text)

    chunks = []
    for block in re.split(r"\n\s*\n", text):
        lines = [
            l.strip()
            for l in block.splitlines()
            if l.strip() and not re.fullmatch(r"\d{1,3}", l.strip())
        ]
        if not lines:
            continue

        joined = re.sub(r"\s+", " ", " ".join(lines)).strip()
        joined = _strip_tail_noise(joined)

        if not joined:
            continue

        if _SECTION_TITLE_RE.match(joined):
            continue

        chunks.append(joined)

    units: List[str] = []
    for chunk in chunks:
        sents = [s.strip() for s in _SENT_SPLIT.split(chunk) if s.strip()]
        if not sents:
            sents = [chunk]

        for sent in sents:
            sent = _strip_tail_noise(sent)

            if len(sent.split()) > 130:
                sub = [
                    s.strip()
                    for s in re.split(
                        r"\s+(?=(?:Khoảng|Vào khoảng|Vào lúc|Hồi|Đến khoảng|Ngày|Tại|Đối với|Trong đó)\b)",
                        sent,
                    )
                    if s.strip()
                ]
                units.extend(sub if sub else [sent])
            else:
                units.append(sent)

    return units


def _is_strict_crime_pattern(unit: str) -> bool:
    return (
        _has_crime_action(unit)
        and _has_time_signal(unit)
        and _has_loc_signal(unit)
        and _has_drug_or_weight(unit)
    )


def _is_drug_name_keep(unit: str) -> bool:
    if not _has_drug_name(unit):
        return False

    n = _strip_diacritics(unit)

    if re.search(r"(NUOC\s+TIEU|MAU\s+NUOC\s+TIEU|TEST\s+NHANH|XET\s+NGHIEM\s+DUONG\s+TINH|DANH\s+MUC|NGHI\s+DINH|STT)", n):
        return False

    if re.search(r"(HOAN\s+LAI|CON\s+LAI\s+SAU\s+GIAM\s+DINH|MAU\s+VAT\s+CON\s+LAI|MAU\s+HOAN\s+LAI)", n):
        return False

    return True


def _is_relevant_content_unit(unit: str) -> bool:
    if not unit or len(unit.split()) < 4:
        return False

    if _is_noise(unit):
        return False

    if _is_strict_crime_pattern(unit):
        return True

    if _is_drug_name_keep(unit):
        return True

    return False


def _dedupe_units(units: List[str]) -> List[str]:
    kept: List[str] = []
    seen: set[str] = set()
    for unit in units:
        key = re.sub(r"\s+", " ", unit.strip())
        if key and key not in seen:
            kept.append(key)
            seen.add(key)
    return kept


_NEGATIVE_KEEP_PAT = re.compile(r"""
    \b(?:
        nơi\s+cư\s+trú
      | địa\s+chỉ\s+cư\s+trú
      | chỗ\s+ở
      | hộ\s+khẩu
      | HKTT
      | trú\s+tại
      | thường\s+trú
      | tạm\s+trú
      | sinh\s+năm
      | con\s+ông
      | con\s+bà
      | Tòa\s+án\s+nhân\s+dân
      | trụ\s+sở\s+Tòa\s+án
      | xét\s+xử\s+sơ\s+thẩm
      | xét\s+xử\s+phúc\s+thẩm
      | phiên\s+tòa
      | ngày\s+xét\s+xử
      | thụ\s+lý
      | Viện\s+kiểm\s+sát
      | bản\s+cáo\s+trạng
      | cáo\s+trạng\s+số
      | truy\s+tố
      | Cơ\s+quan\s+Cảnh\s+sát\s+điều\s+tra
      | Cơ\s+quan\s+điều\s+tra
      | Công\s+an
      | Phòng\s+kỹ\s+thuật\s+hình\s+sự
      | Kết\s+luận\s+giám\s+định
      | kết\s+luận\s+giám\s+định\s+số
      | ngày\s+giám\s+định
      | mẫu\s+nước\s+tiểu
      | nước\s+tiểu
      | test\s+nhanh
      | xét\s+nghiệm\s+dương\s+tính
      | hoàn\s+lại\s+sau\s+giám\s+định
      | còn\s+lại\s+sau\s+giám\s+định
      | mẫu\s+vật\s+còn\s+lại
      | mẫu\s+hoàn\s+lại
    )\b
""", re.VERBOSE | re.IGNORECASE)


_NEGATIVE_DROP_PAT = re.compile(r"""
    \b(?:
        án\s+phí
      | quyền\s+kháng\s+cáo
      | xử\s+lý\s+vật\s+chứng
      | vật\s+chứng
      | tịch\s+thu
      | tiêu\s+hủy
      | sung\s+công
      | nơi\s+nhận
      | thi\s+hành\s+án
      | Nghị\s+quyết\s+số\s*326
      | Điều\s+106\s+Bộ\s+luật\s+Tố\s+tụng
    )\b
""", re.VERBOSE | re.IGNORECASE)


def _is_useful_negative_unit(unit: str) -> bool:
    unit = unit.strip()
    if not unit or len(unit.split()) < 4:
        return False

    if _SECTION_TITLE_RE.match(unit):
        return False

    if _NEGATIVE_DROP_PAT.search(unit):
        return False

    if not _NEGATIVE_KEEP_PAT.search(unit):
        return False

    if _is_strict_crime_pattern(unit):
        return False

    return True


def split_noi_dung_train_units(text: str) -> Tuple[str, str]:
    positive: List[str] = []
    negative: List[str] = []

    for unit in _to_units(text):
        unit = _strip_tail_noise(unit.strip())
        if not unit:
            continue

        if _is_relevant_content_unit(unit):
            positive.append(unit)
        elif _is_useful_negative_unit(unit):
            negative.append(unit)

    return "\n\n".join(_dedupe_units(positive)).strip(), "\n\n".join(_dedupe_units(negative)).strip()


def keep_noi_dung_raw(text: str) -> str:
    positive, _ = split_noi_dung_train_units(text)
    return positive


_ND_HARD_STOP_PAT = re.compile(r"""
    ^\s*\[?\d{0,2}\]?\s*
    (?:
        Về\s+tình\s+tiết\s+tăng\s+nặng
      | Về\s+tình\s+tiết\s+giảm\s+nhẹ
      | Tình\s+tiết\s+tăng\s+nặng
      | Tình\s+tiết\s+giảm\s+nhẹ
      | Về\s+nhân\s+thân
      | Về\s+hình\s+phạt
      | Xét\s+tính\s+chất
      | Về\s+vật\s+chứng
      | Xử\s+lý\s+vật\s+chứng
      | Về\s+án\s+phí
      | Về\s+quyền\s+kháng\s+cáo
      | Các\s+vấn\s+đề\s+khác
      | Vì\s+các\s+lẽ\s+trên
    )\b
""", re.VERBOSE | re.IGNORECASE)

_ND_PROSECUTION_PAT = re.compile(r"""
    \b(?:
        Viện\s+kiểm\s+sát
      | Kiểm\s+sát\s+viên
      | đại\s+diện\s+Viện\s+kiểm\s+sát
      | bản\s+cáo\s+trạng
      | cáo\s+trạng\s+số
      | truy\s+tố
      | luận\s+tội
      | đề\s+nghị\s+Hội\s+đồng\s+xét\s+xử
      | đề\s+nghị\s+xử\s+phạt
    )\b
""", re.VERBOSE | re.IGNORECASE)

_ND_FACT_CUE_PAT = re.compile(r"""
    \b(?:
        có\s+đủ\s+căn\s+cứ\s+kết\s+luận
      | đủ\s+cơ\s+sở\s+kết\s+luận
      | Hội\s+đồng\s+xét\s+xử\s+có\s+đủ\s+cơ\s+sở\s+kết\s+luận
      | Tòa\s+án\s+có\s+đủ\s+cơ\s+sở\s+kết\s+luận
      | Vào\s+khoảng
      | Vào\s+lúc
      | Hồi
      | Khoảng
      | bắt\s+quả\s+tang
      | bị\s+bắt\s+quả\s+tang
      | phát\s+hiện
      | thu\s+giữ
      | tàng\s+trữ\s+trái\s+phép
      | mua\s+bán\s+trái\s+phép
      | vận\s+chuyển\s+trái\s+phép
      | sử\s+dụng\s+trái\s+phép
      | tổ\s+chức\s+sử\s+dụng
      | đã\s+có\s+hành\s+vi
      | hành\s+vi\s+nêu\s+trên
      | phạm\s+vào\s+tội
    )\b
""", re.VERBOSE | re.IGNORECASE)

_ND_DROP_INLINE_PAT = re.compile(r"""
    \b(?:
        án\s+phí
      | vật\s+chứng
      | tịch\s+thu
      | tiêu\s+hủy
      | sung\s+công
      | quyền\s+kháng\s+cáo
      | nơi\s+nhận
      | thi\s+hành\s+án
      | Nghị\s+quyết\s+số\s*326
      | Điều\s+106\s+Bộ\s+luật\s+Tố\s+tụng
      | tình\s+tiết\s+tăng\s+nặng
      | tình\s+tiết\s+giảm\s+nhẹ
      | thành\s+khẩn\s+khai\s+báo
      | ăn\s+năn\s+hối\s+cải
      | mức\s+án
      | hình\s+phạt\s+bổ\s+sung
      | miễn\s+án\s+phí
      | buộc\s+bị\s+cáo\s+phải\s+nộp
      | biết\s+rõ
      | gây\s+tác\s+hại
      | nghiêm\s+cấm
      | xử\s+lý\s+rất\s+nghiêm
      | cần\s+phải\s+xử\s+lý
      | răn\s+đe
      | phòng\s+ngừa\s+chung
      | xâm\s+phạm
      | trật\s+tự\s+công\s+cộng
    )\b
""", re.VERBOSE | re.IGNORECASE)


def _is_reasoning_fact_unit(unit: str) -> bool:
    unit = unit.strip()
    if not unit or len(unit.split()) < 4:
        return False

    if _SECTION_TITLE_RE.match(unit):
        return False

    if _ND_HARD_STOP_PAT.search(unit):
        return False

    has_conclusion = bool(re.search(r"(có\s+đủ\s+căn\s+cứ\s+kết\s+luận|đủ\s+cơ\s+sở\s+kết\s+luận|đã\s+có\s+hành\s+vi|hành\s+vi\s+nêu\s+trên)", unit, re.IGNORECASE))
    if _ND_PROSECUTION_PAT.search(unit) and not has_conclusion:
        return False

    if _ND_DROP_INLINE_PAT.search(unit) and not has_conclusion:
        return False

    n = _strip_diacritics(unit)
    if re.search(r"(MAU\s+NUOC\s+TIEU|NUOC\s+TIEU|TEST\s+NHANH|XET\s+NGHIEM|DANH\s+MUC|NGHI\s+DINH|STT)", n) and not has_conclusion:
        return False

    if has_conclusion:
        return True

    if _is_strict_crime_pattern(unit):
        return True

    if _ND_FACT_CUE_PAT.search(unit) and (_has_drug_or_weight(unit) or _has_time_signal(unit) or _has_loc_signal(unit)):
        return True

    if _has_crime_action(unit) and (_has_drug_or_weight(unit) or (_has_time_signal(unit) and _has_loc_signal(unit))):
        return True

    return False


def filter_nhan_dinh_crime_facts(text: str) -> str:
    kept: List[str] = []
    seen: set[str] = set()

    for unit in _to_units(text):
        unit = _strip_tail_noise(unit.strip())
        if not _is_reasoning_fact_unit(unit):
            continue

        key = re.sub(r"\s+", " ", unit)
        if key and key not in seen:
            kept.append(key)
            seen.add(key)

    return "\n\n".join(kept).strip()


def split_nhan_dinh_train_units(text: str) -> Tuple[str, str]:
    positive_text = filter_nhan_dinh_crime_facts(text)
    positive_set = {
        re.sub(r"\s+", " ", x.strip())
        for x in positive_text.split("\n\n")
        if x.strip()
    }

    negative: List[str] = []
    for unit in _to_units(text):
        unit = _strip_tail_noise(unit.strip())
        key = re.sub(r"\s+", " ", unit)
        if not key or key in positive_set:
            continue
        if _is_useful_negative_unit(unit):
            negative.append(unit)

    return positive_text, "\n\n".join(_dedupe_units(negative)).strip()


def filter_content_or_reasoning_section(text: str) -> str:
    return filter_nhan_dinh_crime_facts(text)


def filter_drug_section(text: str) -> str:
    return filter_content_or_reasoning_section(text)


def filter_crime_fact_section(text: str) -> str:
    return filter_content_or_reasoning_section(text)


def filter_quyet_dinh(text: str) -> str:
    lines = [
        l.strip()
        for l in text.splitlines()
        if l.strip() and not re.fullmatch(r"\d{1,3}", l.strip())
    ]
    merged = re.sub(r"\s+", " ", " ".join(lines)).strip()

    norm = _strip_diacritics(merged)
    stop = len(merged)

    for pat in _QD_STOP:
        m = re.search(pat, norm, re.IGNORECASE | re.MULTILINE)
        if m and m.start() < stop:
            stop = m.start()

    kept = merged[:stop].strip()
    kept = re.sub(r"^QUYẾT\s+ĐỊNH\s*:\s*", "", kept, flags=re.IGNORECASE).strip()

    parts = re.split(r"(?=\b\d{1,2}\.\s+[A-ZĐÁÀẢÃẠ])", kept)

    return "\n\n".join(p.strip() for p in parts if p.strip())


def extract_ner_sections(full_text: str) -> Tuple[str, str, str, str]:
    text = clean_raw(full_text)
    content_pos, nd_pos, qd_pos = _locate_sections(text)

    if content_pos is not None and nd_pos is not None and nd_pos > content_pos:
        content_raw = text[content_pos:nd_pos]
    elif content_pos is not None:
        content_raw = text[content_pos:qd_pos or len(text)]
    else:
        content_raw = ""

    if nd_pos is not None:
        nd_raw = text[nd_pos:qd_pos or len(text)]
    else:
        nd_raw = ""

    qd_raw = text[qd_pos:] if qd_pos is not None else ""

    noi_dung_pos, noi_dung_neg = split_noi_dung_train_units(content_raw) if content_raw else ("", "")
    nhan_dinh_pos, nhan_dinh_neg = split_nhan_dinh_train_units(nd_raw) if nd_raw else ("", "")

    negative_parts = [x for x in [noi_dung_neg, nhan_dinh_neg] if x.strip()]
    negative_context = "\n\n".join(negative_parts).strip()

    quyet_dinh = filter_quyet_dinh(qd_raw) if qd_raw else ""

    return noi_dung_pos, nhan_dinh_pos, negative_context, quyet_dinh


def save_docx(output_path: Path, noi_dung: str, nhan_dinh: str, negative_context: str, quyet_dinh: str) -> None:
    doc = Document()
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(13)

    if noi_dung:
        doc.add_heading("NỘI DUNG VỤ ÁN - ĐOẠN GÁN NHÃN", level=2)
        for para in noi_dung.split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())

    if nhan_dinh:
        doc.add_heading("NHẬN ĐỊNH CỦA TÒA ÁN - ĐOẠN GÁN NHÃN", level=2)
        for para in nhan_dinh.split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())

    if negative_context:
        doc.add_heading("NGỮ CẢNH KHÔNG GÁN NHÃN", level=2)
        for para in negative_context.split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())

    if quyet_dinh:
        doc.add_heading("QUYẾT ĐỊNH", level=2)
        for para in quyet_dinh.split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def get_new_output_path_or_skip(output_dir: Path, filename: str) -> Optional[Path]:
    main_path = output_dir / filename

    if main_path.exists():
        return None

    new_dir = output_dir / "new"
    new_dir.mkdir(parents=True, exist_ok=True)

    candidate = new_dir / filename

    if not candidate.exists():
        return candidate

    stem = candidate.stem
    suffix = candidate.suffix

    idx = 1
    while True:
        numbered = new_dir / f"{stem}_{idx}{suffix}"
        if not numbered.exists():
            return numbered
        idx += 1

def process_pdf(pdf_path: Path, output_dir: Path) -> Optional[Path]:
    stem = re.sub(r"\(\d+\)$", "", pdf_path.stem).strip()
    filename = f"{stem}.docx"

    out_docx = get_new_output_path_or_skip(output_dir, filename)

    if out_docx is None:
        print(f"[SKIP] {pdf_path.name} → đã có trong {output_dir / filename}")
        return None

    full_text = extract_pdf(pdf_path)
    noi_dung, nhan_dinh, negative_context, quyet_dinh = extract_ner_sections(full_text)

    save_docx(out_docx, noi_dung, nhan_dinh, negative_context, quyet_dinh)

    return out_docx


def iter_pdfs(input_path: Path) -> List[Path]:
    if input_path.is_file() and input_path.suffix.lower() == ".pdf":
        return [input_path]
    return sorted(input_path.glob("*.pdf"))


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Trích DOCX NER ma túy: positive snippets + negative context, bỏ án phí/kháng cáo/vật chứng."
    )
    parser.add_argument(
        "--input",
        "-i",
        type=Path,
        default=DEFAULT_INPUT_DIR,
        help="File PDF hoặc thư mục chứa PDF.",
    )
    parser.add_argument(
        "--output",
        "-o",
        type=Path,
        default=DEFAULT_OUTPUT_DIR,
        help="Thư mục xuất DOCX.",
    )
    args = parser.parse_args()

    pdfs = iter_pdfs(args.input)
    if not pdfs:
        raise SystemExit(f"Không tìm thấy PDF trong: {args.input}")

    ok, failed = 0, []

    for pdf in pdfs:
        try:
            docx_p = process_pdf(pdf, args.output)

            if docx_p is None:
                continue

            print(f"[OK]   {pdf.name}")
            print(f"       DOCX → {docx_p}")
            ok += 1
        except Exception as e:
            print(f"[FAIL] {pdf.name}: {e}")
            failed.append((pdf.name, str(e)))

    print(f"\nHoàn tất: {ok}/{len(pdfs)} file thành công.")
    if failed:
        print("Các file lỗi:")
        for name, err in failed:
            print(f"  - {name}: {err}")


if __name__ == "__main__":
    main()
